from langchain_core.documents import Document
import pandas as pd
from pypdf import PdfReader
import os

import os
import pandas as pd
from langchain_core.documents import Document
from pypdf import PdfReader
from docx import Document as DocxDocument

def load_pdfs(folder_path):
    documents = []

    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".pdf"):
                file_path = os.path.join(root, file)

                try:
                    reader = PdfReader(file_path)
                    text = ""

                    for page in reader.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text += extracted + "\n"

                    if text.strip():
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": file,
                                    "path": file_path,
                                    "type": "pdf"
                                }
                            )
                        )

                except Exception as e:
                    print(f"Error reading PDF {file}: {e}")

    return documents

def load_docx(folder_path):
    documents = []

    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(".docx"):
                file_path = os.path.join(root, file)

                try:
                    doc = DocxDocument(file_path)
                    text = "\n".join([p.text for p in doc.paragraphs])

                    if text.strip():
                        documents.append(
                            Document(
                                page_content=text,
                                metadata={
                                    "source": file,
                                    "path": file_path,
                                    "type": "docx"
                                }
                            )
                        )

                except Exception as e:
                    print(f"Error reading DOCX {file}: {e}")

    return documents

def load_structured_data(folder_path):
    documents = []

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)

            try:
                if file.lower().endswith(".csv"):
                    df = pd.read_csv(file_path)

                elif file.lower().endswith((".xlsx", ".xls")):
                    df = pd.read_excel(file_path)

                else:
                    continue

                text = df.to_string(index=False)

                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": file,
                            "path": file_path,
                            "type": "structured"
                        }
                    )
                )

            except Exception as e:
                print(f"Error reading structured file {file}: {e}")

    return documents

def load_all_data(base_path):
    documents = []

    print("Loading data...")

    # Evidence
    documents += load_pdfs(os.path.join(base_path, "evidence"))
    documents += load_docx(os.path.join(base_path, "evidence"))

    # Candidate pack
    documents += load_pdfs(os.path.join(base_path, "candidate_pack"))
    documents += load_docx(os.path.join(base_path, "candidate_pack"))

    # Policies
    documents += load_pdfs(os.path.join(base_path, "policies"))
    documents += load_docx(os.path.join(base_path, "policies"))

    # Reference
    documents += load_pdfs(os.path.join(base_path, "reference"))
    documents += load_docx(os.path.join(base_path, "reference"))

    # Structured data
    documents += load_structured_data(os.path.join(base_path, "structured"))

    print(f"Total documents loaded: {len(documents)}")

    return documents